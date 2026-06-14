from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import json
import io
import re
import traceback
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import copy
from datetime import datetime

app = Flask(__name__)
CORS(app)

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_el(ns, tag):
    return etree.Element(f'{{{ns}}}{tag}')

def sub_el(parent, ns, tag):
    return etree.SubElement(parent, f'{{{ns}}}{tag}')

def make_run(text, italic=True, bold=False):
    r = make_el(MATH_NS, 'r')
    rpr = sub_el(r, MATH_NS, 'rPr')
    if not italic:
        sty = sub_el(rpr, MATH_NS, 'sty')
        sty.set(f'{{{MATH_NS}}}val', 'b' if bold else 'p')
    else:
        sty = sub_el(rpr, MATH_NS, 'sty')
        sty.set(f'{{{MATH_NS}}}val', 'bi' if bold else 'i')
    wrpr = sub_el(r, W_NS, 'rPr')
    rfonts = sub_el(wrpr, W_NS, 'rFonts')
    rfonts.set(f'{{{W_NS}}}ascii', 'Cambria Math')
    rfonts.set(f'{{{W_NS}}}hAnsi', 'Cambria Math')
    t = sub_el(r, MATH_NS, 't')
    t.text = text
    t.set(f'{{{W_NS}}}space', 'preserve')
    return r

def make_text_run(text):
    return make_run(text, italic=False, bold=False)

def make_frac(num_elements, den_elements):
    f = make_el(MATH_NS, 'f')
    fpr = sub_el(f, MATH_NS, 'fPr')
    ftype = sub_el(fpr, MATH_NS, 'type')
    ftype.set(f'{{{MATH_NS}}}val', 'bar')
    num = sub_el(f, MATH_NS, 'num')
    for el in num_elements:
        num.append(el)
    den = sub_el(f, MATH_NS, 'den')
    for el in den_elements:
        den.append(el)
    return f

def make_sup(base_elements, sup_elements):
    ssup = make_el(MATH_NS, 'sSup')
    e = sub_el(ssup, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    s = sub_el(ssup, MATH_NS, 'sup')
    for el in sup_elements:
        s.append(el)
    return ssup

def make_sub_el(base_elements, sub_elements):
    ssub = make_el(MATH_NS, 'sSub')
    e = sub_el(ssub, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    s = sub_el(ssub, MATH_NS, 'sub')
    for el in sub_elements:
        s.append(el)
    return ssub

def make_subsup(base_elements, sub_elements, sup_elements):
    ssubsup = make_el(MATH_NS, 'sSubSup')
    e = sub_el(ssubsup, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    sb = sub_el(ssubsup, MATH_NS, 'sub')
    for el in sub_elements:
        sb.append(el)
    sp = sub_el(ssubsup, MATH_NS, 'sup')
    for el in sup_elements:
        sp.append(el)
    return ssubsup

def make_sqrt(content_elements, degree_elements=None):
    rad = make_el(MATH_NS, 'rad')
    radpr = sub_el(rad, MATH_NS, 'radPr')
    if degree_elements is None:
        deghide = sub_el(radpr, MATH_NS, 'degHide')
        deghide.set(f'{{{MATH_NS}}}val', '1')
    deg = sub_el(rad, MATH_NS, 'deg')
    if degree_elements:
        for el in degree_elements:
            deg.append(el)
    e = sub_el(rad, MATH_NS, 'e')
    for el in content_elements:
        e.append(el)
    return rad

def make_accent(base_elements, accent_char='\u0302'):
    acc = make_el(MATH_NS, 'acc')
    accpr = sub_el(acc, MATH_NS, 'accPr')
    ch = sub_el(accpr, MATH_NS, 'chr')
    ch.set(f'{{{MATH_NS}}}val', accent_char)
    e = sub_el(acc, MATH_NS, 'e')
    for el in base_elements:
        e.append(el)
    return acc

def make_delim(content_elements, beg='(', end=')'):
    d = make_el(MATH_NS, 'd')
    dpr = sub_el(d, MATH_NS, 'dPr')
    begchr = sub_el(dpr, MATH_NS, 'begChr')
    begchr.set(f'{{{MATH_NS}}}val', beg)
    endchr = sub_el(dpr, MATH_NS, 'endChr')
    endchr.set(f'{{{MATH_NS}}}val', end)
    e = sub_el(d, MATH_NS, 'e')
    for el in content_elements:
        e.append(el)
    return d

def make_func(func_name, arg_elements):
    func = make_el(MATH_NS, 'func')
    funcpr = sub_el(func, MATH_NS, 'funcPr')
    fname = sub_el(func, MATH_NS, 'fName')
    fname.append(make_run(func_name, italic=False))
    e = sub_el(func, MATH_NS, 'e')
    for el in arg_elements:
        e.append(el)
    return func

def make_nary(symbol, sub_els=None, sup_els=None, content_els=None):
    nary = make_el(MATH_NS, 'nary')
    narypr = sub_el(nary, MATH_NS, 'naryPr')
    ch = sub_el(narypr, MATH_NS, 'chr')
    ch.set(f'{{{MATH_NS}}}val', symbol)
    if sub_els is None:
        subhide = sub_el(narypr, MATH_NS, 'subHide')
        subhide.set(f'{{{MATH_NS}}}val', '1')
    if sup_els is None:
        suphide = sub_el(narypr, MATH_NS, 'supHide')
        suphide.set(f'{{{MATH_NS}}}val', '1')
    sb = sub_el(nary, MATH_NS, 'sub')
    if sub_els:
        for el in sub_els:
            sb.append(el)
    sp = sub_el(nary, MATH_NS, 'sup')
    if sup_els:
        for el in sup_els:
            sp.append(el)
    e = sub_el(nary, MATH_NS, 'e')
    if content_els:
        for el in content_els:
            e.append(el)
    return nary


GREEK = {
    r'\alpha': 'α', r'\beta': 'β', r'\gamma': 'γ', r'\delta': 'δ',
    r'\epsilon': 'ε', r'\varepsilon': 'ε', r'\zeta': 'ζ', r'\eta': 'η',
    r'\theta': 'θ', r'\vartheta': 'ϑ', r'\iota': 'ι', r'\kappa': 'κ',
    r'\lambda': 'λ', r'\mu': 'μ', r'\nu': 'ν', r'\xi': 'ξ',
    r'\pi': 'π', r'\rho': 'ρ', r'\sigma': 'σ', r'\tau': 'τ',
    r'\upsilon': 'υ', r'\phi': 'φ', r'\varphi': 'φ', r'\chi': 'χ',
    r'\psi': 'ψ', r'\omega': 'ω',
    r'\Gamma': 'Γ', r'\Delta': 'Δ', r'\Theta': 'Θ', r'\Lambda': 'Λ',
    r'\Xi': 'Ξ', r'\Pi': 'Π', r'\Sigma': 'Σ', r'\Upsilon': 'Υ',
    r'\Phi': 'Φ', r'\Psi': 'Ψ', r'\Omega': 'Ω',
}

SYMBOLS = {
    r'\hbar': 'ℏ', r'\infty': '∞', r'\partial': '∂',
    r'\nabla': '∇', r'\pm': '±', r'\mp': '∓',
    r'\times': '×', r'\cdot': '·', r'\cdots': '⋯', r'\ldots': '…',
    r'\leq': '≤', r'\geq': '≥', r'\le': '≤', r'\ge': '≥',
    r'\neq': '≠', r'\ne': '≠', r'\approx': '≈', r'\equiv': '≡',
    r'\sim': '∼', r'\simeq': '≃', r'\propto': '∝',
    r'\rightarrow': '→', r'\leftarrow': '←', r'\Rightarrow': '⇒',
    r'\Leftarrow': '⇐', r'\leftrightarrow': '↔', r'\to': '→',
    r'\forall': '∀', r'\exists': '∃', r'\in': '∈', r'\notin': '∉',
    r'\subset': '⊂', r'\supset': '⊃', r'\subseteq': '⊆', r'\supseteq': '⊇',
    r'\cup': '∪', r'\cap': '∩', r'\emptyset': '∅',
    r'\circ': '∘', r'\bullet': '•', r'\star': '⋆',
    r'\prime': '′', r'\angle': '∠', r'\perp': '⊥', r'\parallel': '∥',
    r'\quad': ' ', r'\qquad': '  ',
}

FUNCTIONS = {
    r'\sin', r'\cos', r'\tan', r'\cot', r'\sec', r'\csc',
    r'\arcsin', r'\arccos', r'\arctan',
    r'\sinh', r'\cosh', r'\tanh', r'\coth',
    r'\ln', r'\log', r'\exp', r'\lim', r'\min', r'\max',
    r'\det', r'\dim', r'\ker', r'\deg',
    r'\arg', r'\sup', r'\inf', r'\gcd',
}


def parse_latex(latex):
    elements = []
    i = 0
    s = latex.strip()
    max_iterations = len(s) * 3 + 100
    iterations = 0

    while i < len(s):
        iterations += 1
        if iterations > max_iterations:
            elements.append(make_run(s[i:], italic=False))
            break

        c = s[i]

        if c == ' ':
            i += 1
            continue

        if c == '{':
            depth = 1
            j = i + 1
            while j < len(s) and depth > 0:
                if s[j] == '{': depth += 1
                elif s[j] == '}': depth -= 1
                j += 1
            inner = s[i+1:j-1]
            elements.extend(parse_latex(inner))
            i = j
            continue

        if c == '\\':
            j = i + 1
            if j < len(s) and not s[j].isalpha():
                special = s[i:j+1]
                if special in (r'\ ', r'\,', r'\;', r'\!', r'\:'):
                    elements.append(make_run(' ', italic=False))
                elif special == r'\\':
                    pass
                else:
                    elements.append(make_run(s[j], italic=False))
                i = j + 1
                continue

            while j < len(s) and s[j].isalpha():
                j += 1
            cmd = s[i:j]

            if cmd == r'\frac':
                num_c, after_n = _read_group(s, j)
                den_c, after_d = _read_group(s, after_n)
                num_els = parse_latex(num_c) or [make_run(' ')]
                den_els = parse_latex(den_c) or [make_run(' ')]
                elements.append(make_frac(num_els, den_els))
                i = after_d
                continue

            if cmd in (r'\text', r'\mathrm', r'\textrm', r'\textbf', r'\operatorname'):
                content, after = _read_group(s, j)
                bold = (cmd == r'\textbf')
                elements.append(make_run(content, italic=False, bold=bold))
                i = after
                continue

            accents_map = {
                r'\hat': '\u0302', r'\vec': '\u20D7', r'\bar': '\u0305',
                r'\tilde': '\u0303', r'\dot': '\u0307', r'\ddot': '\u0308',
                r'\overline': '\u0305', r'\underline': '\u0332',
                r'\widehat': '\u0302', r'\widetilde': '\u0303',
            }
            if cmd in accents_map:
                content, after = _read_group(s, j)
                inner_els = parse_latex(content) or [make_run(' ')]
                elements.append(make_accent(inner_els, accents_map[cmd]))
                i = after
                continue

            if cmd == r'\sqrt':
                deg_els = None
                pos = j
                while pos < len(s) and s[pos] == ' ':
                    pos += 1
                if pos < len(s) and s[pos] == '[':
                    end_bracket = s.find(']', pos)
                    if end_bracket > 0:
                        deg_content = s[pos+1:end_bracket]
                        deg_els = parse_latex(deg_content)
                        pos = end_bracket + 1
                content, after = _read_group(s, pos)
                inner_els = parse_latex(content) or [make_run(' ')]
                elements.append(make_sqrt(inner_els, deg_els))
                i = after
                continue

            if cmd == r'\left':
                beg_char = s[j] if j < len(s) else '('
                if beg_char == '.': beg_char = ''
                right_pos = _find_matching_right(s, j+1)
                if right_pos >= 0:
                    inner = s[j+1:right_pos]
                    end_pos = right_pos + 6
                    end_char = s[end_pos] if end_pos < len(s) else ')'
                    if end_char == '.': end_char = ''
                    inner_els = parse_latex(inner) or [make_run(' ')]
                    elements.append(make_delim(inner_els, beg_char or '(', end_char or ')'))
                    i = end_pos + 1
                else:
                    elements.append(make_run(beg_char, italic=False))
                    i = j + 1
                continue

            if cmd == r'\right':
                i = j + 1
                continue

            if cmd in FUNCTIONS:
                func_name = cmd[1:]
                elements.append(make_run(func_name, italic=False))
                i = j
                continue

            nary_map = {r'\sum': '∑', r'\prod': '∏', r'\int': '∫',
                        r'\iint': '∬', r'\iiint': '∭', r'\oint': '∮'}
            if cmd in nary_map:
                elements.append(make_run(nary_map[cmd], italic=False))
                i = j
                continue

            if cmd in GREEK:
                elements.append(make_run(GREEK[cmd], italic=True))
                i = j
                continue

            if cmd in SYMBOLS:
                elements.append(make_run(SYMBOLS[cmd], italic=False))
                i = j
                continue

            if cmd in (r'\mathbf', r'\mathbb', r'\mathcal', r'\boldsymbol'):
                content, after = _read_group(s, j)
                is_bold = cmd in (r'\mathbf', r'\boldsymbol')
                elements.append(make_run(content, italic=False, bold=is_bold))
                i = after
                continue

            elements.append(make_run(cmd[1:], italic=False))
            i = j
            continue

        if c == '^':
            sup_c, after = _read_group_or_char(s, i+1)
            sup_els = parse_latex(sup_c) or [make_run(' ')]
            if elements:
                base = elements.pop()
                if after < len(s) and s[after] == '_':
                    sub_c, after2 = _read_group_or_char(s, after+1)
                    sub_els2 = parse_latex(sub_c) or [make_run(' ')]
                    elements.append(make_subsup([base], sub_els2, sup_els))
                    i = after2
                else:
                    elements.append(make_sup([base], sup_els))
                    i = after
            else:
                elements.append(make_sup([make_run(' ')], sup_els))
                i = after
            continue

        if c == '_':
            sub_c, after = _read_group_or_char(s, i+1)
            sub_els = parse_latex(sub_c) or [make_run(' ')]
            if elements:
                base = elements.pop()
                if after < len(s) and s[after] == '^':
                    sup_c2, after2 = _read_group_or_char(s, after+1)
                    sup_els2 = parse_latex(sup_c2) or [make_run(' ')]
                    elements.append(make_subsup([base], sub_els, sup_els2))
                    i = after2
                else:
                    elements.append(make_sub_el([base], sub_els))
                    i = after
            else:
                elements.append(make_sub_el([make_run(' ')], sub_els))
                i = after
            continue

        text = ''
        while i < len(s) and s[i] not in '\\{}^_$ \t':
            ch = s[i]
            if ch in '+-=<>':
                if text:
                    elements.append(make_run(text))
                    text = ''
                elements.append(make_run(ch, italic=False))
                i += 1
                continue
            if ch in '(),.:;!?[]|/':
                if text:
                    elements.append(make_run(text))
                    text = ''
                elements.append(make_run(ch, italic=False))
                i += 1
                continue
            text += ch
            i += 1
        if text:
            elements.append(make_run(text))
        continue

    return elements


def _find_matching_right(s, start):
    depth = 1
    i = start
    while i < len(s) - 5:
        if s[i:i+5] == r'\left':
            depth += 1
            i += 5
        elif s[i:i+6] == r'\right':
            depth -= 1
            if depth == 0:
                return i
            i += 6
        else:
            i += 1
    return -1


def _read_group(s, pos):
    while pos < len(s) and s[pos] == ' ':
        pos += 1
    if pos >= len(s):
        return ('', pos)
    if s[pos] == '{':
        depth = 1
        j = pos + 1
        while j < len(s) and depth > 0:
            if s[j] == '{': depth += 1
            elif s[j] == '}': depth -= 1
            j += 1
        return (s[pos+1:j-1], j)
    else:
        return (s[pos], pos+1)


def _read_group_or_char(s, pos):
    while pos < len(s) and s[pos] == ' ':
        pos += 1
    if pos >= len(s):
        return ('', pos)
    if s[pos] == '{':
        return _read_group(s, pos)
    elif s[pos] == '\\':
        j = pos + 1
        while j < len(s) and s[j].isalpha():
            j += 1
        return (s[pos:j], j)
    else:
        return (s[pos], pos+1)









def make_matrix(env, rows_data):
    beg_map = {'pmatrix': '(', 'bmatrix': '[', 'vmatrix': '|', 'matrix': ''}
    end_map = {'pmatrix': ')', 'bmatrix': ']', 'vmatrix': '|', 'matrix': ''}
    beg = beg_map.get(env, '(')
    end = end_map.get(env, ')')
    m = make_el(MATH_NS, 'm')
    mpr = sub_el(m, MATH_NS, 'mPr')
    ncols = max(len(r) for r in rows_data) if rows_data else 1
    mcs = sub_el(mpr, MATH_NS, 'mcs')
    mc_el = sub_el(mcs, MATH_NS, 'mc')
    mcpr = sub_el(mc_el, MATH_NS, 'mcPr')
    count_el = sub_el(mcpr, MATH_NS, 'count')
    count_el.set(f'{{{MATH_NS}}}val', str(ncols))
    mcjc = sub_el(mcpr, MATH_NS, 'mcJc')
    mcjc.set(f'{{{MATH_NS}}}val', 'center')
    for row in rows_data:
        mr = sub_el(m, MATH_NS, 'mr')
        for cell_latex in row:
            e = sub_el(mr, MATH_NS, 'e')
            for el in (parse_latex(cell_latex.strip()) or [make_run(' ')]):
                e.append(el)
        for _ in range(ncols - len(row)):
            e = sub_el(mr, MATH_NS, 'e')
            e.append(make_run(' '))
    if beg and end:
        return make_delim([m], beg, end)
    return m


def parse_matrix_env(latex):
    s = latex.strip()
    bs = chr(92)
    if not s.startswith(bs + 'begin{'):
        return None
    brace_end = s.find('}', 7)
    if brace_end < 0:
        return None
    env = s[7:brace_end]
    if env not in ('pmatrix', 'bmatrix', 'vmatrix', 'matrix', 'smallmatrix'):
        return None
    end_tag = bs + 'end{' + env + '}'
    body_start = brace_end + 1
    body_end = s.find(end_tag)
    if body_end < 0:
        return None
    body = s[body_start:body_end]
    # rows separated by \\ or newlines
    double_bs = bs + bs
    if double_bs in body:
        raw_rows = body.split(double_bs)
    else:
        # fallback: split by newline (when content came from multiline $$ block)
        raw_rows = body.split('\n')
    rows_data = []
    for row in raw_rows:
        row = row.strip()
        if not row:
            continue
        cells = row.split('&')
        rows_data.append([c.strip() for c in cells])
    if not rows_data:
        return None
    return make_matrix(env, rows_data)


def build_omath(latex):
    omath = make_el(MATH_NS, 'oMath')
    s = latex.strip()

    # Full expression is a matrix
    mat = parse_matrix_env(s)
    if mat is not None:
        omath.append(mat)
        return omath

    # Expression contains \\begin{ somewhere (prefix + matrix)
    begin_idx = s.find(chr(92) + 'begin{')
    if begin_idx >= 0:
        prefix = s[:begin_idx].strip()
        rest = s[begin_idx:]
        if prefix:
            for el in (parse_latex(prefix) or []):
                omath.append(el)
        mat2 = parse_matrix_env(rest)
        if mat2 is not None:
            omath.append(mat2)
        else:
            for el in (parse_latex(rest) or []):
                omath.append(el)
        return omath

    elements = parse_latex(latex)
    for el in elements:
        omath.append(el)
    return omath


def insert_math(paragraph, latex):
    try:
        omath = build_omath(latex)
        paragraph._element.append(omath)
        return True
    except Exception as e:
        print(f'Math error "{latex}": {e}')
        r = paragraph.add_run(latex)
        r.font.name = 'Cambria Math'
        r.italic = True
        return False


def add_block_formula(doc, latex):
    """Add a block (display) formula. Handles multiple formulas separated by \\quad or \\,."""
    # Replace \quad and \qquad with spaces for splitting display formulas
    # But first handle matrix environments which may contain these
    latex = latex.strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_math(p, latex)


def process_content(doc, content):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        img = re.search(r'!\[([^\]]*)\]\(([^\)]+)\)', line)
        if img:
            _img(doc, img.group(2), img.group(1))
            i += 1
            continue

        if line.strip().startswith('```'):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code.append(lines[i])
                i += 1
            _code(doc, '\n'.join(code))
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            tlines = []
            while i < len(lines) and '|' in lines[i]:
                if not re.match(r'^\s*\|[\s\-:|]+\|\s*$', lines[i]):
                    tlines.append(lines[i])
                i += 1
            if tlines:
                _table_with_math(doc, tlines)
            continue

        bm = re.match(r'^\s*\$\$(.+?)\$\$\s*$', line)
        if bm:
            add_block_formula(doc, bm.group(1).strip())
            i += 1
            continue

        if line.strip() == '$$':
            fl = []
            i += 1
            while i < len(lines) and lines[i].strip() != '$$':
                fl.append(lines[i])
                i += 1
            raw = '\n'.join(fl).strip()
            if raw:
                # If contains matrix environment, keep newlines; otherwise join with space
                if '\\begin{' in raw:
                    # Convert newlines between matrix rows to \\\\ if not already there
                    latex = raw
                else:
                    latex = ' '.join(fl).strip()
                add_block_formula(doc, latex)
            i += 1
            continue

        if line.strip():
            _text_math(doc, line)
        else:
            doc.add_paragraph()
        i += 1


def _text_math(doc, text):
    parts = re.split(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', text)
    if len(parts) <= 1:
        p = doc.add_paragraph()
        _fmt(p, text)
        return
    p = doc.add_paragraph()
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            if part:
                _fmt(p, part)
        else:
            insert_math(p, part.strip())


def _add_cell_content_with_math(cell, text):
    text = text.strip()
    parts = re.split(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', text)
    para = cell.paragraphs[0]
    if len(parts) <= 1:
        run = para.add_run(text)
        run.font.size = Pt(11)
        return
    for idx, part in enumerate(parts):
        if idx % 2 == 0:
            if part.strip():
                run = para.add_run(part)
                run.font.size = Pt(11)
        else:
            insert_math(para, part.strip())


def _table_with_math(doc, tlines):
    rows = []
    for l in tlines:
        cells = [c.strip() for c in l.split('|')]
        cells = [c for c in cells if c != '']
        if cells:
            rows.append(cells)
    if not rows:
        return
    mc = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=mc)
    t.style = 'Table Grid'
    for i, rd in enumerate(rows):
        for j, ct in enumerate(rd):
            if j < mc:
                cell = t.rows[i].cells[j]
                if '$' in ct:
                    _add_cell_content_with_math(cell, ct)
                else:
                    run = cell.paragraphs[0].add_run(ct)
                    run.font.size = Pt(11)
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def _has_emoji(text):
    import unicodedata
    for ch in text:
        cp = ord(ch)
        if (0x1F300 <= cp <= 0x1FAFF) or (0x2600 <= cp <= 0x27BF) or (0x2300 <= cp <= 0x23FF):
            return True
    return False


def _add_run_with_emoji(para, text, bold=False, italic=False):
    """Add run with Segoe UI Emoji font if text contains emoji."""
    if not text:
        return
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if _has_emoji(text):
        r.font.name = 'Segoe UI Emoji'


def _fmt(para, text):
    bparts = re.split(r'\*\*(.+?)\*\*', text)
    for i, bp in enumerate(bparts):
        if i % 2 == 0:
            iparts = re.split(r'\*(.+?)\*', bp)
            for j, ip in enumerate(iparts):
                if j % 2 == 0:
                    if ip: _add_run_with_emoji(para, ip)
                else:
                    _add_run_with_emoji(para, ip, italic=True)
        else:
            _add_run_with_emoji(para, bp, bold=True)


def _code(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = 'Courier New'
    r.font.size = Pt(10)
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), 'F5F5F5')
    p._element.get_or_add_pPr().append(s)


def _img(doc, src, alt=''):
    try:
        import urllib.request, base64
        if src.startswith('data:image'):
            b64 = src.split('base64,')[1]
            stream = io.BytesIO(base64.b64decode(b64))
        else:
            req = urllib.request.Request(src, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=30) as resp:
                stream = io.BytesIO(resp.read())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(stream, width=Inches(5.0))
    except:
        p = doc.add_paragraph()
        r = p.add_run(f'[Image: {alt}]')
        r.italic = True


@app.route('/', methods=['GET'])
@app.route('/api/export-chat', methods=['GET'])
def health():
    return jsonify({'status': 'OK', 'version': '5.0-flask'})


@app.route('/api/export-chat', methods=['POST'])
def export_chat():
    try:
        data = request.get_json(force=True)
        messages = data.get('messages', [])
        title = data.get('title', 'Gemini Chat')

        if not messages:
            return jsonify({'error': 'No messages'}), 400

        doc = Document()

        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp = doc.add_paragraph()
        dr = dp.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))
        dr.font.size = Pt(10)
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        for i, msg in enumerate(messages):
            role = msg.get('role', 'user')
            content = msg.get('content', '')

            rp = doc.add_paragraph()
            rr = rp.add_run('You' if role == 'user' else 'Gemini')
            rr.bold = True
            rr.font.size = Pt(14)
            rr.font.color.rgb = RGBColor(33, 150, 243) if role == 'user' else RGBColor(76, 175, 80)

            process_content(doc, content)

            if i < len(messages) - 1:
                sp = doc.add_paragraph()
                sr = sp.add_run('─' * 60)
                sr.font.color.rgb = RGBColor(200, 200, 200)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='gemini-chat.docx'
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500


if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port)
